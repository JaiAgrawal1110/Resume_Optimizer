"""
cv_parser.py
Extracts raw text from an uploaded master CV file (PDF, DOCX, or TXT).
This module does NOT structure the content — it just gets clean raw text.
Structuring into the canonical JSON schema happens in groq_service.py (Phase 2).
"""

import io
from typing import Literal

import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


class UnsupportedFileTypeError(Exception):
    pass


class FileTooLargeError(Exception):
    pass


class EmptyContentError(Exception):
    pass


def get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def validate_upload(filename: str, file_bytes: bytes) -> str:
    """Validates extension and size. Returns the normalized extension."""
    ext = get_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '.{ext}'. Supported types: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"File exceeds the {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB limit."
        )
    if len(file_bytes) == 0:
        raise EmptyContentError("Uploaded file is empty.")
    return ext


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts visible text AND hyperlink URLs (e.g. LinkedIn/GitHub icons that
    are clickable links but have no visible URL text). pdfplumber's
    extract_text() only captures visible text, so links rendered as icons
    or generic anchor text ("LinkedIn", "GitHub") would otherwise be lost.
    We append a labeled "Hyperlinks found in document" block so the
    structuring step (groq_service.py) can match URLs to the right fields.
    """
    text_parts = []
    all_links: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            for hyperlink in getattr(page, "hyperlinks", []):
                uri = hyperlink.get("uri")
                if uri and uri not in all_links:
                    all_links.append(uri)

    full_text = "\n".join(text_parts).strip()

    if all_links:
        links_block = "\n".join(f"- {url}" for url in all_links)
        full_text += f"\n\n[Hyperlinks found in document, in document order — use these to fill linkedin_url, github_url, and project links; match by domain/path]\n{links_block}"

    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables (some resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    # Try utf-8 first, fall back to latin-1 for odd encodings
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1").strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Main entrypoint: validates the upload, dispatches to the right extractor,
    and returns raw text. Raises on unsupported type, oversized file, or
    empty/unextractable content.
    """
    ext = validate_upload(filename, file_bytes)

    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": extract_text_from_txt,
    }

    text = extractors[ext](file_bytes)

    if not text or not text.strip():
        raise EmptyContentError(
            "Could not extract any text from this file. "
            "If it's a scanned/image-based PDF, please upload a text-based version instead."
        )

    return text
